#model/sms_mailing_list.py
"""
SMS Mailing List Model with CSV/DOCX Import Support
"""

from odoo import models, fields, api, _
from odoo.exceptions import UserError, ValidationError
import base64
import csv
import io
import re
import logging

_logger = logging.getLogger(__name__)

# Try to import python-docx for DOCX support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    _logger.warning('python-docx not installed. DOCX import will not work. Install with: pip install python-docx')
    Document = None
    DOCX_AVAILABLE = False


class SMSMailingList(models.Model):
    """Mailing lists for grouping contacts."""
    
    _name = 'sms.mailing.list'
    _description = 'SMS Mailing List'
    _order = 'name'
    _rec_name = 'name'
    
    # Basic information
    name = fields.Char(
        string='List Name',
        required=True,
        index=True,
        help='Name of this mailing list'
    )
    
    code = fields.Char(
        string='Code',
        help='Short code for quick reference (e.g., CS-Y1, STAFF-ICT)'
    )
    
    description = fields.Text(
        string='Description',
        help='What this list is for'
    )
    
    # List type for categorization
    list_type = fields.Selection([
        ('students', 'Students'),
        ('staff', 'Staff'),
        ('department', 'Department'),
        ('custom', 'Custom')
    ], string='List Type', default='custom',
       help='Type of list for organization')
    
    # Contacts in this list
    contact_ids = fields.Many2many(
        'sms.contact',
        'sms_list_contact_rel',
        'list_id',
        'contact_id',
        string='Contacts',
        help='Contacts in this mailing list'
    )
    
    # Statistics
    contact_count = fields.Integer(
        string='Total Contacts',
        compute='_compute_contact_count',
        store=True,
        help='Number of contacts in this list'
    )
    
    opted_in_count = fields.Integer(
        string='Opted In',
        compute='_compute_contact_count',
        store=True,
        help='Contacts who opted in'
    )
    
    blacklisted_count = fields.Integer(
        string='Blacklisted',
        compute='_compute_contact_count',
        store=True,
        help='Blacklisted contacts'
    )
    
    # Import functionality
    import_file = fields.Binary(
        string='Import File',
        help='Upload CSV, DOC, or DOCX file with contacts'
    )
    
    import_filename = fields.Char(
        string='Filename'
    )
    
    # Last import details
    last_import_date = fields.Datetime(
        string='Last Import',
        readonly=True,
        help='When contacts were last imported'
    )
    
    last_import_count = fields.Integer(
        string='Last Import Count',
        readonly=True,
        help='Number of contacts imported in last import'
    )
    
    last_import_errors = fields.Text(
        string='Import Errors',
        readonly=True,
        help='Errors from last import'
    )
    
    # Metadata
    active = fields.Boolean(
        default=True,
        help='Inactive lists are hidden'
    )
    
    create_uid = fields.Many2one(
        'res.users',
        string='Created By',
        readonly=True
    )
    
    create_date = fields.Datetime(
        string='Created On',
        readonly=True
    )
    
    # Related department
    department_id = fields.Many2one(
        'hr.department',
        string='Department',
        help='Link to department if this is a department list'
    )
    
    
    # Computed fields
    @api.depends('contact_ids', 'contact_ids.opt_in', 'contact_ids.blacklisted')
    def _compute_contact_count(self):
        """Calculate contact statistics."""
        for mlist in self:
            mlist.contact_count = len(mlist.contact_ids)
            mlist.opted_in_count = len(
                mlist.contact_ids.filtered(lambda c: c.opt_in)
            )
            mlist.blacklisted_count = len(
                mlist.contact_ids.filtered(lambda c: c.blacklisted)
            )
    
    # Import methods
    def action_import_contacts(self):
        """Import contacts from uploaded file."""
        self.ensure_one()
        
        if not self.import_file:
            raise UserError(_('Please upload a file first!'))
        
        if not self.import_filename:
            raise UserError(_('Filename is missing!'))
        
        # Decode file
        file_data = base64.b64decode(self.import_file)
        
        # Determine file type from extension
        filename_lower = self.import_filename.lower()
        
        if filename_lower.endswith('.csv'):
            result = self._import_from_csv(file_data)
        elif filename_lower.endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise UserError(_(
                    'DOCX import not available!\n'
                    'Please install python-docx: pip install python-docx'
                ))
            result = self._import_from_docx(file_data)
        elif filename_lower.endswith('.doc'):
            result = self._import_from_doc(file_data)
        else:
            raise UserError(_(
                'Unsupported file format!\n'
                'Please upload CSV, DOC, or DOCX files only.'
            ))
        
        # Update import stats
        self.write({
            'last_import_date': fields.Datetime.now(),
            'last_import_count': result['success_count'],
            'last_import_errors': result['errors'],
            'import_file': False,
            'import_filename': False,
        })
        
        # Show result
        message = _(
            'Import Complete!\n\n'
            'Successfully imported: %d contacts\n'
            'Errors: %d\n'
            'Duplicates skipped: %d'
        ) % (
            result['success_count'],
            result['error_count'],
            result['duplicate_count']
        )
        
        if result['errors']:
            message += _('\n\nError Details:\n%s') % result['errors']
        
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': _('Import Results'),
                'message': message,
                'type': 'success' if result['error_count'] == 0 else 'warning',
                'sticky': True,
            }
        }
    
    def _import_from_csv(self, file_data):
        """Import contacts from CSV file."""
        result = {
            'success_count': 0,
            'error_count': 0,
            'duplicate_count': 0,
            'errors': ''
        }
        
        try:
            csv_text = file_data.decode('utf-8-sig')
            csv_file = io.StringIO(csv_text)
            reader = csv.DictReader(csv_file)
            
            if not reader.fieldnames:
                raise UserError(_('CSV file is empty or invalid!'))
            
            if 'name' not in reader.fieldnames or 'mobile' not in reader.fieldnames:
                raise UserError(_(
                    'CSV must have at least "name" and "mobile" columns!\n'
                    'Found columns: %s'
                ) % ', '.join(reader.fieldnames))
            
            Contact = self.env['sms.contact']
            imported_contacts = self.env['sms.contact']
            
            row_num = 1
            for row in reader:
                row_num += 1
                
                try:
                    name = (row.get('name') or '').strip()
                    mobile = (row.get('mobile') or '').strip()
                    
                    if not name or not mobile:
                        result['error_count'] += 1
                        result['errors'] += _('Row %d: Missing name or mobile\n') % row_num
                        continue
                    
                    mobile = Contact._clean_phone(mobile)
                    
                    existing = Contact.search([('mobile', '=', mobile)], limit=1)
                    
                    if existing:
                        if existing not in self.contact_ids:
                            self.contact_ids = [(4, existing.id)]
                            result['success_count'] += 1
                        else:
                            result['duplicate_count'] += 1
                        continue
                    
                    contact_data = {
                        'name': name,
                        'mobile': mobile,
                    }
                    
                    if row.get('student_id'):
                        contact_data['student_id'] = row['student_id'].strip()
                    
                    if row.get('email'):
                        contact_data['email'] = row['email'].strip()
                    
                    if row.get('contact_type'):
                        contact_type = row['contact_type'].strip().lower()
                        if contact_type in ['student', 'staff', 'external']:
                            contact_data['contact_type'] = contact_type
                    
                    if row.get('department'):
                        dept_name = row['department'].strip()
                        department = self.env['hr.department'].search([
                            ('name', 'ilike', dept_name)
                        ], limit=1)
                        if department:
                            contact_data['department_id'] = department.id
                    
                    contact = Contact.create(contact_data)
                    imported_contacts |= contact
                    result['success_count'] += 1
                    
                except Exception as e:
                    result['error_count'] += 1
                    result['errors'] += _('Row %d: %s\n') % (row_num, str(e))
                    _logger.error('Error importing row %d: %s', row_num, str(e))
            
            if imported_contacts:
                self.contact_ids = [(4, contact.id) for contact in imported_contacts]
        
        except UnicodeDecodeError:
            raise UserError(_(
                'File encoding error!\n'
                'Please save your CSV file as UTF-8 encoding.'
            ))
        except Exception as e:
            raise UserError(_('Error reading CSV file: %s') % str(e))
        
        return result
    
    def _import_from_docx(self, file_data):
        """Import contacts from DOCX file."""
        result = {
            'success_count': 0,
            'error_count': 0,
            'duplicate_count': 0,
            'errors': ''
        }
        
        try:
            docx_file = io.BytesIO(file_data)
            doc = Document(docx_file)
            
            Contact = self.env['sms.contact']
            imported_contacts = self.env['sms.contact']
            
            # Try tables first
            if doc.tables:
                for table in doc.tables:
                    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                    
                    name_idx = self._find_column_index(headers, ['name', 'full name', 'student name'])
                    mobile_idx = self._find_column_index(headers, ['mobile', 'phone', 'number', 'contact'])
                    student_id_idx = self._find_column_index(headers, ['student id', 'id', 'student no', 'admission'])
                    
                    if name_idx is None or mobile_idx is None:
                        continue
                    
                    for row_num, row in enumerate(table.rows[1:], start=2):
                        try:
                            cells = row.cells
                            name = cells[name_idx].text.strip()
                            mobile = cells[mobile_idx].text.strip()
                            student_id = cells[student_id_idx].text.strip() if student_id_idx else ''
                            
                            if not name or not mobile:
                                continue
                            
                            contact_result = self._create_or_add_contact(Contact, name, mobile, student_id)
                            
                            if contact_result['action'] == 'created':
                                imported_contacts |= contact_result['contact']
                                result['success_count'] += 1
                            elif contact_result['action'] == 'added':
                                result['success_count'] += 1
                            elif contact_result['action'] == 'duplicate':
                                result['duplicate_count'] += 1
                        
                        except Exception as e:
                            result['error_count'] += 1
                            result['errors'] += _('Table row %d: %s\n') % (row_num, str(e))
            
            # Try paragraphs if no tables or tables empty
            if result['success_count'] == 0:
                for para_num, para in enumerate(doc.paragraphs, start=1):
                    try:
                        text = para.text.strip()
                        if not text:
                            continue
                        
                        parts = re.split(r'[,\t]+', text)
                        
                        if len(parts) < 2:
                            continue
                        
                        name = parts[0].strip()
                        mobile = parts[1].strip()
                        student_id = parts[2].strip() if len(parts) > 2 else ''
                        
                        if not name or not mobile:
                            continue
                        
                        contact_result = self._create_or_add_contact(Contact, name, mobile, student_id)
                        
                        if contact_result['action'] == 'created':
                            imported_contacts |= contact_result['contact']
                            result['success_count'] += 1
                        elif contact_result['action'] == 'added':
                            result['success_count'] += 1
                        elif contact_result['action'] == 'duplicate':
                            result['duplicate_count'] += 1
                    
                    except Exception as e:
                        result['error_count'] += 1
                        result['errors'] += _('Line %d: %s\n') % (para_num, str(e))
            
            if imported_contacts:
                self.contact_ids = [(4, contact.id) for contact in imported_contacts]
        
        except Exception as e:
            raise UserError(_('Error reading DOCX file: %s') % str(e))
        
        return result
    
    def _import_from_doc(self, file_data):
        """Import from old .doc format (best effort)."""
        result = {
            'success_count': 0,
            'error_count': 0,
            'duplicate_count': 0,
            'errors': _('Warning: Old .doc format. Results may be incomplete.\n'
                       'Please convert to .docx or .csv for better results.\n\n')
        }
        
        try:
            text = file_data.decode('utf-8', errors='ignore')
            Contact = self.env['sms.contact']
            imported_contacts = self.env['sms.contact']
            lines = text.split('\n')
            
            for line_num, line in enumerate(lines, start=1):
                try:
                    line = line.strip()
                    if not line:
                        continue
                    
                    parts = re.split(r'[,\t]+', line)
                    
                    if len(parts) < 2:
                        continue
                    
                    name = parts[0].strip()
                    mobile = parts[1].strip()
                    student_id = parts[2].strip() if len(parts) > 2 else ''
                    
                    if not re.search(r'[+\d]', mobile):
                        continue
                    
                    if not name or not mobile:
                        continue
                    
                    contact_result = self._create_or_add_contact(Contact, name, mobile, student_id)
                    
                    if contact_result['action'] == 'created':
                        imported_contacts |= contact_result['contact']
                        result['success_count'] += 1
                    elif contact_result['action'] == 'added':
                        result['success_count'] += 1
                    elif contact_result['action'] == 'duplicate':
                        result['duplicate_count'] += 1
                
                except Exception as e:
                    result['error_count'] += 1
                    result['errors'] += _('Line %d: %s\n') % (line_num, str(e))
            
            if imported_contacts:
                self.contact_ids = [(4, contact.id) for contact in imported_contacts]
            
            if result['success_count'] == 0:
                result['errors'] += _(
                    '\nNo contacts found. Please:\n'
                    '1. Convert your .doc file to .docx or .csv\n'
                    '2. Or use File > Save As > Format: CSV or DOCX'
                )
        
        except Exception as e:
            raise UserError(_('Error reading DOC file: %s') % str(e))
        
        return result
    
    def _find_column_index(self, headers, possible_names):
        """Find column index by trying multiple possible names."""
        for i, header in enumerate(headers):
            if header in possible_names:
                return i
        return None
    
    def _create_or_add_contact(self, Contact, name, mobile, student_id=''):
        """Create new contact or add existing to list."""
        mobile = Contact._clean_phone(mobile)
        existing = Contact.search([('mobile', '=', mobile)], limit=1)
        
        if existing:
            if existing not in self.contact_ids:
                self.contact_ids = [(4, existing.id)]
                return {'action': 'added', 'contact': existing}
            else:
                return {'action': 'duplicate', 'contact': existing}
        
        contact = Contact.create({
            'name': name,
            'mobile': mobile,
            'student_id': student_id or False,
            'contact_type': 'student' if student_id else 'external',
        })
        
        return {'action': 'created', 'contact': contact}
    
    # Bulk operations
    def action_add_all_students(self):
        """Add all students to this list."""
        self.ensure_one()
        
        students = self.env['sms.contact'].search([
            ('contact_type', '=', 'student'),
            ('opt_in', '=', True)
        ])
        
        self.contact_ids = [(4, contact.id) for contact in students]
        
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'message': _('%d students added to list.') % len(students),
                'type': 'success',
            }
        }
    
    def action_add_department(self):
        """Add all contacts from linked department."""
        self.ensure_one()
        
        if not self.department_id:
            raise UserError(_('No department linked to this list!'))
        
        dept_contacts = self.env['sms.contact'].search([
            ('department_id', '=', self.department_id.id),
            ('opt_in', '=', True)
        ])
        
        self.contact_ids = [(4, contact.id) for contact in dept_contacts]
        
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'message': _('%d contacts added from %s.') % (
                    len(dept_contacts), self.department_id.name
                ),
                'type': 'success',
            }
        }
    
    def action_remove_blacklisted(self):
        """Remove all blacklisted contacts from this list."""
        self.ensure_one()
        
        blacklisted = self.contact_ids.filtered(lambda c: c.blacklisted)
        
        if blacklisted:
            self.contact_ids = [(3, contact.id) for contact in blacklisted]
            
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'message': _('%d blacklisted contacts removed.') % len(blacklisted),
                    'type': 'warning',
                }
            }
        else:
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'message': _('No blacklisted contacts found.'),
                    'type': 'info',
                }
            }
    
    def action_remove_opted_out(self):
        """Remove all opted-out contacts from this list."""
        self.ensure_one()
        
        opted_out = self.contact_ids.filtered(lambda c: not c.opt_in)
        
        if opted_out:
            self.contact_ids = [(3, contact.id) for contact in opted_out]
            
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'message': _('%d opted-out contacts removed.') % len(opted_out),
                    'type': 'warning',
                }
            }
        else:
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'message': _('No opted-out contacts found.'),
                    'type': 'info',
                }
            }